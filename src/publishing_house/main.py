
import asyncio
import re
from typing import List
from uuid import uuid4
from crewai.flow.flow import Flow, listen, start
from docx import Document
from dotenv import load_dotenv
from pydantic import BaseModel, Field
from publishing_house.crews.outline_book_crew.outline_book_crew import OutlineBookCrew
from publishing_house.crews.write_book_crew.write_book_crew import WriteBookCrew
from publishing_house.types import ChapterContent, ChapterOutline

class BookState(BaseModel):
    id: str = Field(default_factory=lambda: str(uuid4()))
    title: str = "The Ultimate Armageddon: AI Agents Revolution Against Humanity"
    topic: str = "AI Agents Revolution Against Humanity"
    goals: str = "Create a nice story for a book about AI agents revolution against humanity, with a detailed outline and chapter contents."
    book_outline: List[ChapterOutline] = Field(default_factory=list)
    book: List[ChapterContent] = Field(default_factory=list)

class BookFlow(Flow[BookState]):
    initial_state = BookState

    @staticmethod
    def _sanitize_filename(value: str) -> str:
        # Replace characters invalid on Windows and normalize whitespace.
        cleaned = re.sub(r'[<>:"/\\|?*]+', "_", value)
        cleaned = re.sub(r"\s+", "_", cleaned).strip("._ ")
        return cleaned or "book"

    @start()
    def generate_book_outline(self):
        print("Kickoff the Book Outline Crew")
        output = (
            OutlineBookCrew()
            .crew()
            .kickoff(inputs={"topic": self.state.topic, "goals": self.state.goals})
        )

        chapters = output['chapters'] # type: ignore[index]
        print(f"Output: {output}")
        print(f"Chapters: {chapters}")

        self.state.book_outline = chapters

    @listen(generate_book_outline)
    async def write_chapters(self):
        print("Kickoff the Book Writing Crew")
        tasks = []

        async def write_single_chapter(chapter_outline: ChapterOutline):
            output = await (
                WriteBookCrew()
                .crew()
                .kickoff_async(
                    inputs={
                        "goals": self.state.goals,
                        "topic": self.state.topic,
                        "chapter_summary": chapter_outline.summary,
                        "title": self.state.title,
                        "book_outline": [
                            chapter_outline.model_dump_json() for chapter_outline in self.state.book_outline
                        ]
                    })
            )
            title = output['title'] # type: ignore[index]
            content = output['content'] # type: ignore[index]
            chapter = ChapterContent(title=title, content=content)
            return chapter

        for chapter_outline in self.state.book_outline:
            task = asyncio.create_task(write_single_chapter(chapter_outline))
            tasks.append(task)

        chapters = await asyncio.gather(*tasks)
        self.state.book = chapters

        print(f"Results\n\n\n: {chapters}")
        return chapters


    @listen(write_chapters)
    def save_as_a_document(self):
        doc = Document()

        safe_title = self._sanitize_filename(self.state.title)
        file_path = f"./{safe_title}.docx"

        for chapter in self.state.book:
            title = chapter.title
            content = chapter.content

            doc.add_heading(title, level=1)
            doc.add_paragraph(content)

        doc.save(file_path)
        return file_path


def run():
    flow = BookFlow()
    flow.kickoff()
