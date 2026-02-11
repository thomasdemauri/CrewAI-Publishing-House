from crewai.tools import BaseTool
from typing import Type
from pydantic import BaseModel, Field
from docx import Document
import os

from publishing_house.types import DocxWriterInput

class DocxWriterTool(BaseTool):
    name: str = "docx_writer_tool"
    description: str = (
        "Creates a real .docx file from a formatted book manuscript text. "
        "Use this tool to export the final book into a professional Word document."
    )
    args_schema: Type[BaseModel] = DocxWriterInput

    def _run(self, file_path: str, title: str, content: str) -> str:
        try:
            os.makedirs(os.path.dirname(file_path), exist_ok=True)

            doc = Document()

            doc.add_heading(title, level=0)

            paragraphs = content.split("\n")
            for p in paragraphs:
                p = p.strip()
                if not p:
                    doc.add_paragraph("")
                    continue

                if p.lower().startswith("chapter"):
                    doc.add_heading(p, level=1)
                elif p.isupper() and len(p) < 80:
                    doc.add_heading(p, level=2)
                else:
                    doc.add_paragraph(p)

            doc.save(file_path)

            return f"DOCX file successfully generated at: {file_path}"

        except Exception as e:
            return f"Error generating DOCX file: {str(e)}"
